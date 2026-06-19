"""
fem_core.py — Finite Element Method engine
Supports:
  1. Continuous Beam  (Euler-Bernoulli beam elements, DOF: v, θ per node)
  2. Plane Frame      (2-D frame elements, DOF: u, v, θ per node)

Mesh density: n_elem per span = max(2, round(span / (0.01 * total_length)))
All units: kN, m, kNm  (EI in kN·m², EA in kN)
"""
from __future__ import annotations

import io
from dataclasses import dataclass, field
from typing import Optional
import numpy as np


# ══════════════════════════════════════════════════════
#  DATA CLASSES — INPUT
# ══════════════════════════════════════════════════════

@dataclass
class SpanDef:
    """One span of a continuous beam."""
    length: float          # m
    EI: float = 1.0        # kN·m²  (use 1.0 for unit EI diagrams)
    # Loads on this span
    point_loads:    list[tuple[float, float]] = field(default_factory=list)  # (P kN, x_local m)
    point_moments:  list[tuple[float, float]] = field(default_factory=list)  # (M kNm, x_local m)
    udls:           list[tuple[float, float, float]] = field(default_factory=list)  # (q, x1_local, x2_local)

@dataclass
class SupportDef:
    """Support at a node position of the continuous beam."""
    node: int              # 0-based node index
    kind: str = "pin"      # "pin" | "fixed" | "roller" | "free"

@dataclass
class ContinuousBeamInput:
    spans: list[SpanDef]
    supports: list[SupportDef]

# ── Plane Frame ──────────────────────────────────────

@dataclass
class FrameNode:
    x: float
    y: float

@dataclass
class FrameElement:
    i_node: int            # start node index (0-based)
    j_node: int            # end node index
    E: float  = 200e6     # kN/m²  (steel default ≈ 200 GPa)
    A: float  = 0.01      # m²
    I: float  = 1e-4      # m⁴
    # Distributed load in LOCAL y-direction (perpendicular to element axis)
    udl_local: float = 0.0  # kN/m

@dataclass
class FrameSupport:
    node: int
    ux_fixed: bool = True
    uy_fixed: bool = True
    rz_fixed: bool = True  # moment fixity

@dataclass
class FramePointLoad:
    node: int
    Fx: float = 0.0   # kN  (global x)
    Fy: float = 0.0   # kN  (global y)
    Mz: float = 0.0   # kNm

@dataclass
class PlaneFrameInput:
    nodes:       list[FrameNode]
    elements:    list[FrameElement]
    supports:    list[FrameSupport]
    point_loads: list[FramePointLoad] = field(default_factory=list)


# ══════════════════════════════════════════════════════
#  DATA CLASSES — OUTPUT
# ══════════════════════════════════════════════════════

@dataclass
class ContinuousBeamResult:
    x_global: np.ndarray      # global x coordinates
    shear:    np.ndarray
    moment:   np.ndarray
    deflection: np.ndarray
    theta:    np.ndarray
    reactions: dict[int, dict]  # node_index -> {Fy, Mz}
    report: str

@dataclass
class FrameElementResult:
    elem_idx: int
    x_local:  np.ndarray     # 0 → L
    axial:    np.ndarray     # kN  (tension +)
    shear:    np.ndarray     # kN
    moment:   np.ndarray     # kNm
    # global coordinates for plotting
    x_coords: np.ndarray
    y_coords: np.ndarray

@dataclass
class PlaneFrameResult:
    node_displacements: np.ndarray   # shape (n_nodes, 3): [ux, uy, rz]
    reactions: list[dict]            # per support
    element_results: list[FrameElementResult]
    report: str


# ══════════════════════════════════════════════════════
#  CONTINUOUS BEAM  — FEM SOLVER
# ══════════════════════════════════════════════════════

def _beam_element_stiffness(EI: float, L: float) -> np.ndarray:
    """4×4 Euler-Bernoulli beam element stiffness matrix.
    DOF order: [v_i, θ_i, v_j, θ_j]
    """
    k = EI / L**3 * np.array([
        [ 12,   6*L,  -12,   6*L],
        [  6*L, 4*L**2, -6*L, 2*L**2],
        [-12,  -6*L,   12,  -6*L],
        [  6*L, 2*L**2, -6*L, 4*L**2],
    ])
    return k


def _consistent_load_udl(q: float, x1: float, x2: float, L_elem: float) -> np.ndarray:
    """Consistent nodal load vector for UDL on part [x1,x2] of element length L_elem.
    Uses numerical integration (Gauss) over the load region.
    Returns [fv_i, fm_i, fv_j, fm_j]
    """
    # Hermite shape functions: N1,N2,N3,N4 as function of ξ ∈ [0,1]
    # ξ = x / L
    def hermite(xi):
        return np.array([
            1 - 3*xi**2 + 2*xi**3,
            L_elem * xi * (1 - xi)**2,
            3*xi**2 - 2*xi**3,
            L_elem * xi**2 * (xi - 1),
        ])

    # Gauss quadrature over [x1/L, x2/L]
    n_pts = 16
    a, b = x1 / L_elem, x2 / L_elem
    xi_pts = np.linspace(a, b, n_pts)
    dxi = (b - a) / (n_pts - 1)
    f = np.zeros(4)
    for i, xi in enumerate(xi_pts):
        w = dxi if (i == 0 or i == n_pts - 1) else dxi  # trapezoidal
        f += q * hermite(xi) * L_elem * w
    # trapezoidal: endpoints weight 0.5
    f = np.zeros(4)
    for i, xi in enumerate(xi_pts):
        w = (dxi / 2) if (i == 0 or i == n_pts - 1) else dxi
        f += q * hermite(xi) * L_elem * w
    return f


def solve_continuous_beam(data: ContinuousBeamInput, pts_per_elem: int = 20) -> ContinuousBeamResult:
    """Assemble global stiffness, apply BCs, solve, recover diagrams."""

    spans = data.spans
    n_spans = len(spans)
    total_L = sum(s.length for s in spans)

    # Node positions (span interfaces)
    node_x = np.concatenate([[0.0], np.cumsum([s.length for s in spans])])
    n_nodes = len(node_x)  # n_spans + 1

    # ── Mesh: subdivide each span into n_elem elements ──
    mesh_elems = []   # list of (EI, L_local, span_idx, x_start_global)
    elem_nodes = []   # (i_node_global, j_node_global)
    global_node_x = [0.0]

    gnode = 0
    for s_idx, span in enumerate(spans):
        n_elem = max(2, round(span.length / (0.01 * total_L)))
        Le = span.length / n_elem
        for e in range(n_elem):
            mesh_elems.append({
                "EI": span.EI,
                "L": Le,
                "span_idx": s_idx,
                "x0_global": node_x[s_idx] + e * Le,
            })
            elem_nodes.append((gnode, gnode + 1))
            gnode += 1
            global_node_x.append(node_x[s_idx] + (e + 1) * Le)
    global_node_x = np.array(global_node_x)
    n_gnodes = len(global_node_x)  # gnode + 1
    n_dof = 2 * n_gnodes           # [v, θ] per node

    # ── Map span-end nodes to mesh nodes ──
    # span boundary nodes: node 0 = mesh node 0; node k = after sum of n_elem[0..k-1]
    span_node_map = [0]
    acc = 0
    for s_idx, span in enumerate(spans):
        n_elem = max(2, round(span.length / (0.01 * total_L)))
        acc += n_elem
        span_node_map.append(acc)
    # span_node_map[k] = mesh global-node index at left end of span k (0) and at each support

    # ── Assemble global stiffness ──
    K = np.zeros((n_dof, n_dof))
    F = np.zeros(n_dof)

    for e_idx, (e_data, (ni, nj)) in enumerate(zip(mesh_elems, elem_nodes)):
        EI = e_data["EI"]
        Le = e_data["L"]
        ke = _beam_element_stiffness(EI, Le)
        dofs = [2*ni, 2*ni+1, 2*nj, 2*nj+1]
        for a in range(4):
            for b in range(4):
                K[dofs[a], dofs[b]] += ke[a, b]

    # ── Consistent nodal loads from spans ──
    for s_idx, span in enumerate(spans):
        n_elem = max(2, round(span.length / (0.01 * total_L)))
        Le = span.length / n_elem
        gnode_start = span_node_map[s_idx]

        # UDLs
        for q, x1_sp, x2_sp in span.udls:
            # find which elements the load spans
            for e in range(n_elem):
                x_e_start = e * Le
                x_e_end   = (e + 1) * Le
                a_local = max(x1_sp, x_e_start) - x_e_start
                b_local = min(x2_sp, x_e_end)   - x_e_start
                if b_local <= a_local:
                    continue
                fe = _consistent_load_udl(q, a_local, b_local, Le)
                ni = gnode_start + e
                nj = gnode_start + e + 1
                dofs = [2*ni, 2*ni+1, 2*nj, 2*nj+1]
                for k in range(4):
                    F[dofs[k]] += fe[k]

        # Point loads (placed at nearest mesh node within span)
        for P, x_loc in span.point_loads:
            # find element
            e_idx_local = min(int(x_loc / Le), n_elem - 1)
            xi = (x_loc - e_idx_local * Le) / Le
            # Hermite shape functions
            N = np.array([
                1 - 3*xi**2 + 2*xi**3,
                Le * xi * (1 - xi)**2,
                3*xi**2 - 2*xi**3,
                Le * xi**2 * (xi - 1),
            ])
            ni = gnode_start + e_idx_local
            nj = gnode_start + e_idx_local + 1
            dofs = [2*ni, 2*ni+1, 2*nj, 2*nj+1]
            for k in range(4):
                F[dofs[k]] += P * N[k]

        # Point moments
        for M, x_loc in span.point_moments:
            e_idx_local = min(int(x_loc / Le), n_elem - 1)
            xi = (x_loc - e_idx_local * Le) / Le
            # Derivative of Hermite (for moment)
            dN = np.array([
                (-6*xi + 6*xi**2) / Le,
                1 - 4*xi + 3*xi**2,
                (6*xi - 6*xi**2) / Le,
                -2*xi + 3*xi**2,
            ])
            ni = gnode_start + e_idx_local
            nj = gnode_start + e_idx_local + 1
            dofs = [2*ni, 2*ni+1, 2*nj, 2*nj+1]
            for k in range(4):
                F[dofs[k]] += M * dN[k]

    # ── Boundary conditions ──
    # Map span-level support nodes to mesh global nodes
    # Support node index refers to span-boundary nodes (0..n_spans)
    constrained_dofs = []
    for sup in data.supports:
        mesh_gnode = span_node_map[sup.node]
        v_dof  = 2 * mesh_gnode
        th_dof = 2 * mesh_gnode + 1
        if sup.kind in ("pin", "roller", "fixed"):
            constrained_dofs.append(v_dof)
        if sup.kind == "fixed":
            constrained_dofs.append(th_dof)
    constrained_dofs = list(set(constrained_dofs))

    free_dofs = [d for d in range(n_dof) if d not in constrained_dofs]

    K_ff = K[np.ix_(free_dofs, free_dofs)]
    F_f  = F[free_dofs]

    U_f = np.linalg.solve(K_ff, F_f)

    U = np.zeros(n_dof)
    for i, d in enumerate(free_dofs):
        U[d] = U_f[i]

    # ── Reactions ──
    R_full = K @ U - F
    reactions: dict[int, dict] = {}
    for sup in data.supports:
        mesh_gnode = span_node_map[sup.node]
        v_dof  = 2 * mesh_gnode
        th_dof = 2 * mesh_gnode + 1
        reactions[sup.node] = {
            "x_pos": node_x[sup.node],
            "Fy": -R_full[v_dof] if v_dof in constrained_dofs else 0.0,
            "Mz": -R_full[th_dof] if (sup.kind == "fixed" and th_dof in constrained_dofs) else 0.0,
            "kind": sup.kind,
        }

    # ── Post-process: recover V, M, w along beam ──
    x_out_list, V_list, M_list, w_list, th_list = [], [], [], [], []

    for e_idx, (e_data, (ni, nj)) in enumerate(zip(mesh_elems, elem_nodes)):
        EI = e_data["EI"]
        Le = e_data["L"]
        x0g = e_data["x0_global"]
        s_idx = e_data["span_idx"]

        u_e = U[[2*ni, 2*ni+1, 2*nj, 2*nj+1]]
        xi_arr = np.linspace(0, 1, pts_per_elem + 1)
        # avoid duplicate at interface
        if e_idx < len(mesh_elems) - 1:
            xi_arr = xi_arr[:-1]

        for xi in xi_arr:
            x = x0g + xi * Le
            # Hermite shape functions
            N  = np.array([1-3*xi**2+2*xi**3, Le*xi*(1-xi)**2, 3*xi**2-2*xi**3, Le*xi**2*(xi-1)])
            dN = np.array([(-6*xi+6*xi**2)/Le, 1-4*xi+3*xi**2, (6*xi-6*xi**2)/Le, -2*xi+3*xi**2])
            d2N = np.array([(-6+12*xi)/Le**2, (-4+6*xi)/Le, (6-12*xi)/Le**2, (-2+6*xi)/Le])
            d3N = np.array([12/Le**3, 6/Le**2, -12/Le**3, 6/Le**2])

            w_val  = float(N  @ u_e)
            th_val = float(dN @ u_e)
            M_val  = float(EI * (d2N @ u_e))
            V_val  = float(EI * (d3N @ u_e))

            # add distributed load contribution to shear (from loading)
            for q, x1_sp, x2_sp in data.spans[s_idx].udls:
                x_sp_local = x - node_x[s_idx]
                if x1_sp <= x_sp_local <= x2_sp:
                    pass  # already handled via consistent loads in U

            x_out_list.append(x)
            V_list.append(V_val)
            M_list.append(M_val)
            w_list.append(w_val)
            th_list.append(th_val)

    # add last point
    last_e = mesh_elems[-1]
    ni_last, nj_last = elem_nodes[-1]
    u_last = U[[2*ni_last, 2*ni_last+1, 2*nj_last, 2*nj_last+1]]
    EI_last = last_e["EI"]
    Le_last = last_e["L"]
    d3N_end = np.array([12/Le_last**3, 6/Le_last**2, -12/Le_last**3, 6/Le_last**2])
    d2N_end = np.array([(-6+12)/Le_last**2, (-4+6)/Le_last, (6-12)/Le_last**2, (-2+6)/Le_last])
    N_end   = np.array([0,0,1,0])
    dN_end  = np.array([0,0,0,1/Le_last])  # rough
    x_out_list.append(total_L)
    V_list.append(float(EI_last * (d3N_end @ u_last)))
    M_list.append(float(EI_last * (d2N_end @ u_last)))
    w_list.append(float(N_end @ u_last))
    th_list.append(float(dN_end @ u_last))

    x_arr = np.array(x_out_list)
    V_arr = np.array(V_list)
    M_arr = np.array(M_list)
    w_arr = np.array(w_list)
    th_arr = np.array(th_list)

    # small zero cleanup
    for arr in [V_arr, M_arr, w_arr]:
        arr[np.abs(arr) < 1e-10 * (np.max(np.abs(arr)) + 1e-30)] = 0.0

    report = _build_cb_report(data, reactions, x_arr, V_arr, M_arr, w_arr, span_node_map, n_gnodes, mesh_elems)

    return ContinuousBeamResult(
        x_global=x_arr,
        shear=V_arr,
        moment=M_arr,
        deflection=w_arr,
        theta=th_arr,
        reactions=reactions,
        report=report,
    )


def _build_cb_report(data, reactions, x_arr, V_arr, M_arr, w_arr, span_node_map, n_gnodes, mesh_elems) -> str:
    lines = []
    lines.append("=" * 52)
    lines.append("   THUYẾT MINH TÍNH TOÁN — DẦM LIÊN TỤC (FEM)")
    lines.append("=" * 52)
    lines.append("")
    lines.append("1. THÔNG TIN KẾT CẤU")
    lines.append(f"   Số nhịp       : {len(data.spans)}")
    total_L = sum(s.length for s in data.spans)
    lines.append(f"   Tổng chiều dài: {total_L:.3f} m")
    for i, sp in enumerate(data.spans):
        n_el = max(2, round(sp.length / (0.01 * total_L)))
        lines.append(f"   Nhịp {i+1}: L = {sp.length:.3f} m | EI = {sp.EI:.4g} kN·m² | Số phần tử = {n_el}")
    lines.append("")
    lines.append("2. ĐIỀU KIỆN BIÊN (GỐI ĐỠ)")
    sup_names = {"pin": "Gối khớp (pin)", "roller": "Gối con lăn", "fixed": "Ngàm cứng (fixed)", "free": "Đầu tự do"}
    node_x_list = [0.0] + list(np.cumsum([s.length for s in data.spans]))
    for sup in data.supports:
        lines.append(f"   Node {sup.node} (x = {node_x_list[sup.node]:.3f} m): {sup_names.get(sup.kind, sup.kind)}")
    lines.append("")
    lines.append("3. TẢI TRỌNG")
    for i, sp in enumerate(data.spans):
        x0 = node_x_list[i]
        if sp.point_loads or sp.udls or sp.point_moments:
            lines.append(f"   Nhịp {i+1} (x_global = {x0:.3f} → {x0+sp.length:.3f} m):")
        for P, xl in sp.point_loads:
            lines.append(f"     Tải tập trung: P = {P:.3f} kN tại x_local = {xl:.3f} m (x_global = {x0+xl:.3f} m)")
        for q, x1, x2 in sp.udls:
            lines.append(f"     UDL: q = {q:.3f} kN/m từ x_local = {x1:.3f} → {x2:.3f} m")
        for M, xl in sp.point_moments:
            lines.append(f"     Moment tập trung: M = {M:.3f} kNm tại x_local = {xl:.3f} m")
    lines.append("")
    lines.append("4. PHƯƠNG PHÁP GIẢI — PHẦN TỬ HỮU HẠN (FEM)")
    lines.append("   a. Phần tử: Dầm Euler-Bernoulli")
    lines.append("   b. Bậc tự do mỗi nút: [v (chuyển vị), θ (góc xoay)]")
    lines.append("   c. Hàm nội suy: đa thức Hermite bậc 3")
    lines.append("   d. Ma trận độ cứng phần tử (4×4):")
    lines.append("      [12    6L   -12   6L ]")
    lines.append("      [6L    4L²  -6L   2L²] × EI/L³")
    lines.append("      [-12  -6L    12  -6L ]")
    lines.append("      [6L    2L²  -6L   4L²]")
    lines.append(f"   e. Tổng số nút FEM: {n_gnodes}")
    lines.append(f"   f. Tổng số DOF: {2 * n_gnodes}")
    lines.append("   g. Tải phân bố: vector lực nút tương đương (consistent load)")
    lines.append("   h. Giải hệ: K_ff · U_f = F_f  (phương pháp phân hoạch)")
    lines.append("")
    lines.append("5. PHẢN LỰC GỐI")
    for node_idx, rv in reactions.items():
        lines.append(f"   Node {node_idx} (x = {rv['x_pos']:.3f} m) [{rv['kind']}]:")
        lines.append(f"     Phản lực đứng  Fy = {rv['Fy']:.4f} kN")
        if abs(rv['Mz']) > 1e-8:
            lines.append(f"     Phản lực moment Mz = {rv['Mz']:.4f} kNm")
    # Kiểm tra cân bằng
    total_fy = sum(rv['Fy'] for rv in reactions.values())
    total_ext = 0.0
    for sp in data.spans:
        for P, _ in sp.point_loads: total_ext += P
        for q, x1, x2 in sp.udls: total_ext += q * (x2 - x1)
    lines.append(f"   Kiểm tra ΣFy: Phản lực = {total_fy:.4f} kN | Tải ngoài = {total_ext:.4f} kN")
    if abs(total_fy - total_ext) < 1e-4:
        lines.append("   >> ΣFy = 0 ✓ (Thỏa mãn)")
    else:
        lines.append(f"   >> ΔΣFy = {abs(total_fy - total_ext):.2e} kN (kiểm tra lại điều kiện biên)")
    lines.append("")
    lines.append("6. KẾT QUẢ NỘI LỰC")
    idx_v = int(np.argmax(np.abs(V_arr)))
    idx_m = int(np.argmax(np.abs(M_arr)))
    idx_w = int(np.argmax(np.abs(w_arr)))
    lines.append(f"   Vmax = {V_arr[idx_v]:+.4f} kN    tại x = {x_arr[idx_v]:.3f} m")
    lines.append(f"   Mmax = {M_arr[idx_m]:+.4f} kNm   tại x = {x_arr[idx_m]:.3f} m")
    lines.append(f"   wmax = {w_arr[idx_w]:+.6f} m/EI  tại x = {x_arr[idx_w]:.3f} m")
    lines.append("")
    lines.append("   (EI theo đơn vị nhập; nếu EI = 1.0 thì wmax là chuyển vị trên EI)")
    return "\n".join(lines)


# ══════════════════════════════════════════════════════
#  PLANE FRAME — FEM SOLVER
# ══════════════════════════════════════════════════════

def _frame_element_stiffness_local(E, A, I, L) -> np.ndarray:
    """6×6 local stiffness matrix for 2D frame element.
    DOF order: [u_i, v_i, θ_i, u_j, v_j, θ_j]
    """
    EA_L  = E * A / L
    EI_L3 = E * I / L**3
    EI_L2 = E * I / L**2
    EI_L  = E * I / L

    k = np.array([
        [ EA_L,       0,        0,  -EA_L,       0,        0],
        [    0,  12*EI_L3,  6*EI_L2,     0, -12*EI_L3,  6*EI_L2],
        [    0,   6*EI_L2,  4*EI_L,      0,  -6*EI_L2,  2*EI_L ],
        [-EA_L,       0,        0,   EA_L,       0,        0],
        [    0, -12*EI_L3, -6*EI_L2,    0,  12*EI_L3, -6*EI_L2],
        [    0,   6*EI_L2,  2*EI_L,      0,  -6*EI_L2,  4*EI_L ],
    ])
    return k


def _rotation_matrix(c, s) -> np.ndarray:
    """6×6 transformation matrix T (local → global)."""
    T = np.zeros((6, 6))
    T[0, 0] =  c;  T[0, 1] = s
    T[1, 0] = -s;  T[1, 1] = c
    T[2, 2] =  1
    T[3, 3] =  c;  T[3, 4] = s
    T[4, 3] = -s;  T[4, 4] = c
    T[5, 5] =  1
    return T


def _frame_elem_global_stiffness(elem: FrameElement, nodes: list[FrameNode]):
    ni, nj = nodes[elem.i_node], nodes[elem.j_node]
    dx = nj.x - ni.x
    dy = nj.y - ni.y
    L = np.sqrt(dx**2 + dy**2)
    if L < 1e-12:
        raise ValueError(f"Phần tử {elem.i_node}-{elem.j_node} có chiều dài = 0.")
    c, s = dx / L, dy / L
    k_local = _frame_element_stiffness_local(elem.E, elem.A, elem.I, L)
    T = _rotation_matrix(c, s)
    k_global = T.T @ k_local @ T
    return k_global, T, L, c, s


def _frame_udl_local_to_global(elem: FrameElement, L: float, T: np.ndarray) -> np.ndarray:
    """Consistent nodal loads for udl_local (perpendicular, local y) in global coords."""
    q = elem.udl_local
    fe_local = np.array([
        0,
        q * L / 2,
        q * L**2 / 12,
        0,
        q * L / 2,
        -q * L**2 / 12,
    ])
    fe_global = T.T @ fe_local
    return fe_global


def solve_plane_frame(data: PlaneFrameInput) -> PlaneFrameResult:
    n_nodes = len(data.nodes)
    n_dof   = 3 * n_nodes   # [ux, uy, rz] per node

    K = np.zeros((n_dof, n_dof))
    F = np.zeros(n_dof)

    elem_cache = []  # store (k_global, T, L, c, s) per element

    for e_idx, elem in enumerate(data.elements):
        k_global, T, L, c, s = _frame_elem_global_stiffness(elem, data.nodes)
        elem_cache.append((k_global, T, L, c, s))

        dofs = _frame_elem_dofs(elem)
        for a in range(6):
            for b in range(6):
                K[dofs[a], dofs[b]] += k_global[a, b]

        # UDL
        if elem.udl_local != 0.0:
            fe = _frame_udl_local_to_global(elem, L, T)
            for a in range(6):
                F[dofs[a]] += fe[a]

    # Point loads at nodes
    for pl in data.point_loads:
        d = 3 * pl.node
        F[d]     += pl.Fx
        F[d + 1] += pl.Fy
        F[d + 2] += pl.Mz

    # Boundary conditions
    constrained = []
    for sup in data.supports:
        d = 3 * sup.node
        if sup.ux_fixed: constrained.append(d)
        if sup.uy_fixed: constrained.append(d + 1)
        if sup.rz_fixed: constrained.append(d + 2)
    constrained = list(set(constrained))
    free_dofs = [d for d in range(n_dof) if d not in constrained]

    K_ff = K[np.ix_(free_dofs, free_dofs)]
    F_f  = F[free_dofs]

    try:
        U_f = np.linalg.solve(K_ff, F_f)
    except np.linalg.LinAlgError:
        raise ValueError("Ma trận độ cứng suy biến — kiểm tra lại điều kiện biên (có thể thiếu gối đỡ).")

    U = np.zeros(n_dof)
    for i, d in enumerate(free_dofs):
        U[d] = U_f[i]

    # Reactions
    R_full = K @ U - F
    reactions_out = []
    for sup in data.supports:
        d = 3 * sup.node
        reactions_out.append({
            "node": sup.node,
            "x": data.nodes[sup.node].x,
            "y": data.nodes[sup.node].y,
            "Rx": -R_full[d]     if sup.ux_fixed else 0.0,
            "Ry": -R_full[d + 1] if sup.uy_fixed else 0.0,
            "Mz": -R_full[d + 2] if sup.rz_fixed else 0.0,
        })

    # Element post-processing
    n_disp = len(data.nodes)
    node_disp = U.reshape(n_nodes, 3)

    elem_results = []
    pts = 30
    for e_idx, elem in enumerate(data.elements):
        k_global, T, L, c, s = elem_cache[e_idx]
        dofs = _frame_elem_dofs(elem)
        u_global = U[dofs]
        u_local  = T @ u_global

        # local post-process
        xi_arr = np.linspace(0, 1, pts)
        ax_arr  = np.zeros(pts)
        sh_arr  = np.zeros(pts)
        mo_arr  = np.zeros(pts)
        x_g_arr = np.zeros(pts)
        y_g_arr = np.zeros(pts)

        ni_node = data.nodes[elem.i_node]
        nj_node = data.nodes[elem.j_node]

        EI  = elem.E * elem.I
        EA  = elem.E * elem.A
        q   = elem.udl_local

        for k, xi in enumerate(xi_arr):
            x_loc = xi * L
            # Axial (linear)
            N_ax = np.array([1 - xi, xi])
            axial_val = EA / L * (u_local[3] - u_local[0])
            ax_arr[k] = axial_val

            # Transverse — Hermite
            v_i, th_i, v_j, th_j = u_local[1], u_local[2], u_local[4], u_local[5]
            N_v  = np.array([1-3*xi**2+2*xi**3, L*xi*(1-xi)**2, 3*xi**2-2*xi**3, L*xi**2*(xi-1)])
            d2N  = np.array([(-6+12*xi)/L**2, (-4+6*xi)/L, (6-12*xi)/L**2, (-2+6*xi)/L])
            d3N  = np.array([12/L**3, 6/L**2, -12/L**3, 6/L**2])
            u_bend = np.array([v_i, th_i, v_j, th_j])

            M_val = EI * (d2N @ u_bend) + q * L**2 * xi * (xi - 1) / 2
            V_val = EI * (d3N @ u_bend) + q * L * (0.5 - xi)

            sh_arr[k] = V_val
            mo_arr[k] = M_val

            # Global coords
            x_g_arr[k] = ni_node.x + xi * (nj_node.x - ni_node.x)
            y_g_arr[k] = ni_node.y + xi * (nj_node.y - ni_node.y)

        elem_results.append(FrameElementResult(
            elem_idx=e_idx,
            x_local=xi_arr * L,
            axial=ax_arr,
            shear=sh_arr,
            moment=mo_arr,
            x_coords=x_g_arr,
            y_coords=y_g_arr,
        ))

    report = _build_frame_report(data, node_disp, reactions_out, elem_results)

    return PlaneFrameResult(
        node_displacements=node_disp,
        reactions=reactions_out,
        element_results=elem_results,
        report=report,
    )


def _frame_elem_dofs(elem: FrameElement) -> list[int]:
    i, j = elem.i_node, elem.j_node
    return [3*i, 3*i+1, 3*i+2, 3*j, 3*j+1, 3*j+2]


def _build_frame_report(data: PlaneFrameInput, node_disp, reactions, elem_results) -> str:
    lines = []
    lines.append("=" * 52)
    lines.append("   THUYẾT MINH TÍNH TOÁN — KHUNG PHẲNG (FEM)")
    lines.append("=" * 52)
    lines.append("")
    lines.append("1. THÔNG TIN KẾT CẤU")
    lines.append(f"   Số nút      : {len(data.nodes)}")
    lines.append(f"   Số phần tử  : {len(data.elements)}")
    lines.append(f"   Tổng số DOF : {3 * len(data.nodes)}  ([ux, uy, θz] mỗi nút)")
    lines.append("")
    lines.append("2. TỌA ĐỘ NÚT")
    for i, nd in enumerate(data.nodes):
        lines.append(f"   Nút {i:3d}: x = {nd.x:10.4f} m  |  y = {nd.y:10.4f} m")
    lines.append("")
    lines.append("3. PHẦN TỬ")
    for i, el in enumerate(data.elements):
        ni, nj = data.nodes[el.i_node], data.nodes[el.j_node]
        L = np.sqrt((nj.x-ni.x)**2 + (nj.y-ni.y)**2)
        lines.append(f"   PT {i:3d}: Nút {el.i_node} → {el.j_node} | L = {L:.4f} m | E = {el.E:.4g} kN/m² | A = {el.A:.4g} m² | I = {el.I:.4g} m⁴")
        if el.udl_local != 0:
            lines.append(f"          UDL cục bộ: q = {el.udl_local:.4f} kN/m (phương y cục bộ)")
    lines.append("")
    lines.append("4. TẢI TẬP TRUNG TẠI NÚT")
    if data.point_loads:
        for pl in data.point_loads:
            lines.append(f"   Nút {pl.node}: Fx = {pl.Fx:.4f} kN | Fy = {pl.Fy:.4f} kN | Mz = {pl.Mz:.4f} kNm")
    else:
        lines.append("   (Không có tải tập trung tại nút)")
    lines.append("")
    lines.append("5. ĐIỀU KIỆN BIÊN")
    for sup in data.supports:
        flags = []
        if sup.ux_fixed: flags.append("ux=0")
        if sup.uy_fixed: flags.append("uy=0")
        if sup.rz_fixed: flags.append("θz=0")
        lines.append(f"   Nút {sup.node}: {', '.join(flags)}")
    lines.append("")
    lines.append("6. PHƯƠNG PHÁP GIẢI — FEM KHUNG PHẲNG")
    lines.append("   a. Phần tử: Thanh chịu uốn + nén/kéo đồng thời (Euler-Bernoulli + thanh dàn)")
    lines.append("   b. Ma trận cứng cục bộ 6×6 (DOF: u_i, v_i, θ_i, u_j, v_j, θ_j)")
    lines.append("   c. Chuyển sang tọa độ tổng thể: K_global = Tᵀ · k_local · T")
    lines.append("   d. Lắp ghép ma trận tổng thể K (kích thước 3n × 3n)")
    lines.append("   e. Giải: K_ff · U = F_f  (loại bỏ DOF bị ràng buộc)")
    lines.append("   f. Phục hồi nội lực trong hệ tọa độ cục bộ từng phần tử")
    lines.append("")
    lines.append("7. CHUYỂN VỊ NÚT")
    lines.append(f"   {'Nút':>4}  {'ux (m)':>14}  {'uy (m)':>14}  {'θz (rad)':>14}")
    lines.append("   " + "-" * 52)
    for i, disp in enumerate(node_disp):
        lines.append(f"   {i:>4}  {disp[0]:>14.6e}  {disp[1]:>14.6e}  {disp[2]:>14.6e}")
    lines.append("")
    lines.append("8. PHẢN LỰC GỐI")
    lines.append(f"   {'Nút':>4}  {'Rx (kN)':>12}  {'Ry (kN)':>12}  {'Mz (kNm)':>12}")
    lines.append("   " + "-" * 46)
    for rv in reactions:
        lines.append(f"   {rv['node']:>4}  {rv['Rx']:>12.4f}  {rv['Ry']:>12.4f}  {rv['Mz']:>12.4f}")
    lines.append("")
    # Equilibrium check
    sum_Rx = sum(r['Rx'] for r in reactions)
    sum_Ry = sum(r['Ry'] for r in reactions)
    sum_Mz = sum(r['Mz'] for r in reactions)
    ext_Fx = sum(pl.Fx for pl in data.point_loads)
    ext_Fy = sum(pl.Fy for pl in data.point_loads)
    for el in data.elements:
        ni, nj = data.nodes[el.i_node], data.nodes[el.j_node]
        L = np.sqrt((nj.x-ni.x)**2 + (nj.y-ni.y)**2)
        ext_Fy += el.udl_local * L
    lines.append(f"   Kiểm tra ΣFx: {sum_Rx + ext_Fx:+.4e} kN")
    lines.append(f"   Kiểm tra ΣFy: {sum_Ry + ext_Fy:+.4e} kN")
    lines.append("")
    lines.append("9. KẾT QUẢ NỘI LỰC CÁC PHẦN TỬ")
    for er in elem_results:
        idx_v = int(np.argmax(np.abs(er.shear)))
        idx_m = int(np.argmax(np.abs(er.moment)))
        ax_avg = float(np.mean(er.axial))
        lines.append(f"   Phần tử {er.elem_idx:3d}:")
        lines.append(f"     Lực dọc (trung bình) N  = {ax_avg:+.4f} kN")
        lines.append(f"     Lực cắt lớn nhất    Vmax = {er.shear[idx_v]:+.4f} kN   tại x_local = {er.x_local[idx_v]:.3f} m")
        lines.append(f"     Moment lớn nhất     Mmax = {er.moment[idx_m]:+.4f} kNm  tại x_local = {er.x_local[idx_m]:.3f} m")
    return "\n".join(lines)


# ══════════════════════════════════════════════════════
#  DOCX REPORT BUILDER  (called from web_app.py)
# ══════════════════════════════════════════════════════

def build_docx_bytes(report_text: str, title: str = "Thuyết Minh Tính Toán") -> bytes:
    """
    Generates a .docx file in-memory from plain report text.
    Uses python-docx (pip install python-docx).
    Returns raw bytes of the docx file.
    """
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx chưa được cài. Chạy: pip install python-docx")

    doc = DocxDocument()

    # Page setup
    section = doc.sections[0]
    section.page_width  = int(21 * 360000 / 10)   # A4 width  in EMU (≈ 21cm)
    section.page_height = int(29.7 * 360000 / 10) # A4 height in EMU
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Title
    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x10, 0x18, 0x28)
        run.font.size = Pt(16)

    doc.add_paragraph()

    # Body — split by lines, detect section headers
    for line in report_text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("="):
            continue  # decorative separator
        elif stripped and stripped[0].isdigit() and ". " in stripped[:4]:
            # Section heading
            p = doc.add_heading(stripped, level=2)
            for run in p.runs:
                run.font.size = Pt(12)
        elif stripped.startswith("---") or stripped.startswith("==="):
            doc.add_paragraph()
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()